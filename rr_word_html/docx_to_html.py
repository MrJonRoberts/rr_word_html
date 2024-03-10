from docx import Document
from bs4 import BeautifulSoup
import re

import rr_word_html


class DocxToHtml:

    def __init__(self):
        self._html_path = "./output"
        self._docx_path = "./input"


    def _save_file(self, file_name, content):
        file_path = f"{self._html_path}/{file_name}"
        with open(file_path, "w", encoding="utf-8", errors="replace_with_space") as file:
            file.write(content)
    def _get_document_styles(self, docx_path):
        doc = Document(docx_path)
        styles = set()

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name.replace(" ", "_")
            styles.add(style_name)


        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        styles.add(paragraph.style.name.replace(" ", "_"))

        return styles

    def _generate_css_for_styles(self, styles):
        # basic styles
        file_name = "styles.css"
        # basic reset from https://www.digitalocean.com/community/tutorials/css-minimal-css-reset
        css = rr_word_html.css

        for style in styles:
            css += f".{style.lower()} {{}}\n"
        # for style in styles:
        #     if "Heading" in style:
        #         level = style[-1]
        #         css += f".{style.lower()} {{ font-weight: bold; font-size: {20 - int(level) * 2}px; }}\n"
        #     elif "Normal" in style:
        #         css += f".{style.lower()} {{ font-size: 16px; }}\n"
        #     # Add more mappings based on the styles you find

        self._save_file(file_name, css)
        link = f"<link rel='stylesheet' type='text/css' href='{file_name}'>"

        return link

    def convert_paragraph(self, paragraph):
        style_class = paragraph.style.name.lower()
        if paragraph.style.name.startswith('Heading'):
            level = paragraph.style.name[-1]
            p_html = f'<h{level} class="{style_class}">'
        else:
            p_html = f'<p class="{style_class}">'

        # Process runs as before, omitting direct style application in favor of classes
        for run in paragraph.runs:
            p_html += run.text  # Simplified for demonstration

        p_html += f'</h{level}>' if paragraph.style.name.startswith('Heading') else "</p>"
        return p_html

    def _replace_special_characters(self, error):
        #return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return (" ", error.start + 1)

    def _replace_smart_quotes(self, text):
        # Mapping of smart quotes to straight quotes
        smart_quotes_map = {
            '\u2018': "'",  # Left single quotation mark
            '\u2019': "'",  # Right single quotation mark
            '\u201C': '"',  # Left double quotation mark
            '\u201D': '"',  # Right double quotation mark
        }

        # Replace each smart quote with its straight counterpart
        for smart, straight in smart_quotes_map.items():
            text = text.replace(smart, straight)

        return text

    def _remove_blank_html_elements(self, html_content):
        soup = BeautifulSoup(html_content, 'html.parser')
        ignored_tags = {'link', 'img', 'br'}
        # Find all elements
        for element in soup.find_all(lambda tag: tag.name not in ignored_tags, recursive=True):

            # Check if an element is "empty" (whitespace only or actually empty)
            # Check if an element is considered "blank"
            if (not element.contents or
                    all(isinstance(c, str) and c.isspace() for c in element.contents)):
                element.decompose()  # Remove the element from the parse tree

        return str(soup)

    def convert(self, docx_path):
        doc = Document(docx_path)
        styles = self._get_document_styles(docx_path)
        css = self._generate_css_for_styles(styles)
        print(f">>>CSS: {css}")
        html_output = f"<html><head><title>Document</title>{css}</head><body>"

        for paragraph in doc.paragraphs:
            html_output += self.convert_paragraph(paragraph)

        html_output = self._remove_blank_html_elements(html_output)
        html_output = self._replace_smart_quotes(html_output)

        # html_output += "</body></html>"

        self._save_file("output.html", html_output)
        # with open("output.html", "w", encoding="utf-8", errors="replace_with_space") as file:
        #     file.write(html_output)
